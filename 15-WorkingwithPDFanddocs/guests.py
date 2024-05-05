from docx import Document

guest_name = [ 'Prof. Plum', 'Miss Scarlet', 'Col. Mustard', 'Al Sweigart', 'RoboCop']
doc = Document()

for name in guest_name:
    doc.add_paragraph('It would be a pleasure to have the company of')
    doc.add_paragraph(name)
    doc.add_paragraph('''at 11010 Memory Lane on the Evening of
                               April 1st
                              at 7 o'clock  ''')
    doc.add_page_break()


doc.save('customisedguests.docx')

