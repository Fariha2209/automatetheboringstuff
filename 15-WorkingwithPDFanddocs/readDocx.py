import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '/n'.join(fullText)
    
import readDocx
print(readDocx.getText('/Users/farihaalam/Downloads/automate_online-materials/demo.docx'))
 
doc = docx.Document('/Users/farihaalam/Downloads/automate_online-materials/demo.docx')
doc.paragraphs[0].text

